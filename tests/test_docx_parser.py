import os
import sys
import zipfile
from io import BytesIO

from docx import Document

ROOT_DIR = os.path.dirname(os.path.dirname(__file__))
if ROOT_DIR not in sys.path:
    sys.path.insert(0, ROOT_DIR)

from backend.utils.docx_parser import DocxParser


def _inject_docx_embedding(host_bytes: bytes, embedded_name: str, embedded_bytes: bytes) -> bytes:
    source = BytesIO(host_bytes)
    target = BytesIO()
    with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            zout.writestr(item, zin.read(item.filename))
        zout.writestr(f"word/embeddings/{embedded_name}", embedded_bytes)
    return target.getvalue()


def test_parse_uses_first_non_header_table_cell_for_basic_info(tmp_path):
    document = Document()
    table = document.add_table(rows=3, cols=4)

    name_row = table.rows[0].cells
    name_row[0].text = "需求名称"
    name_row[1].text = "需求名称"
    name_row[2].text = "2026年管理会计集市系统中收模型优化等共计33个迭代优化需求"
    name_row[3].text = "2026年管理会计集市系统中收模型优化等共计33个迭代优化需求"

    summary_row = table.rows[1].cells
    summary_row[0].text = "需求概述"
    summary_row[1].text = "需求概述"
    summary_row[2].text = "针对经营分析与绩效考核要求，管会集市进行系统优化"
    summary_row[3].text = "针对经营分析与绩效考核要求，管会集市进行系统优化"

    content_row = table.rows[2].cells
    content_row[0].text = "需求功能要点描述"
    content_row[1].text = "需求功能要点描述"
    content_row[2].text = "管会集市从取分润补录表改为接入CRM业绩分润数据"
    content_row[3].text = "管会集市从取分润补录表改为接入CRM业绩分润数据"

    file_path = tmp_path / "sample.docx"
    document.save(file_path)

    result = DocxParser().parse(str(file_path))

    assert result["requirement_name"] == "2026年管理会计集市系统中收模型优化等共计33个迭代优化需求"
    assert result["requirement_summary"] == "针对经营分析与绩效考核要求，管会集市进行系统优化"


def test_parse_extracts_requirement_content_from_merged_table_cell(tmp_path):
    document = Document()
    table = document.add_table(rows=2, cols=4)

    for cell in table.rows[0].cells:
        cell.text = "需求内容说明"

    merged = table.rows[1].cells[0].merge(table.rows[1].cells[-1])
    merged.text = (
        "需求功能要点描述\n"
        "押品系统新增配置复核需求\n"
        "新增岗位：配置复核岗；\n"
        "二、押品系统相关优化\n"
        "1.不动产接口字段类型调整。\n"
        "7.押品系统新增DNS改造功能"
    )

    file_path = tmp_path / "merged-table-content.docx"
    document.save(file_path)

    result = DocxParser().parse(str(file_path))

    assert "新增岗位：配置复核岗" in result["requirement_content"]
    assert "二、押品系统相关优化" in result["requirement_content"]
    assert "1.不动产接口字段类型调整" in result["requirement_content"]
    assert "7.押品系统新增DNS改造功能" in result["requirement_content"]
    assert result["requirement_content"].count("押品系统新增配置复核需求") == 1


def test_parse_merges_embedded_attachment_text_into_requirement_content(tmp_path):
    host = Document()
    host.add_paragraph("需求名称")
    host.add_paragraph("测试需求")
    host.add_paragraph("需求内容说明")
    host.add_paragraph("功能描述")
    host.add_paragraph("主文档正文")
    host.add_paragraph("领导审核意见")

    embedded = Document()
    embedded.add_paragraph("附件正文")

    host_buf = BytesIO()
    embedded_buf = BytesIO()
    host.save(host_buf)
    embedded.save(embedded_buf)

    merged_bytes = _inject_docx_embedding(host_buf.getvalue(), "attachment.docx", embedded_buf.getvalue())
    file_path = tmp_path / "merged.docx"
    file_path.write_bytes(merged_bytes)

    result = DocxParser().parse(str(file_path))

    assert "主文档正文" in result["requirement_content"]
    assert "附件正文" in result["requirement_content"]


def test_parse_merges_embedded_txt_attachment_text_into_requirement_content(tmp_path):
    host = Document()
    host.add_paragraph("需求名称")
    host.add_paragraph("测试需求")
    host.add_paragraph("需求内容说明")
    host.add_paragraph("功能描述")
    host.add_paragraph("主文档正文")
    host.add_paragraph("领导审核意见")

    host_buf = BytesIO()
    host.save(host_buf)

    merged_bytes = _inject_docx_embedding(host_buf.getvalue(), "notes.txt", "txt附件正文".encode("utf-8"))
    file_path = tmp_path / "with-txt.docx"
    file_path.write_bytes(merged_bytes)

    result = DocxParser().parse(str(file_path))

    assert "主文档正文" in result["requirement_content"]
    assert "txt附件正文" in result["requirement_content"]
    assert "【附件: notes.txt】" in result["requirement_content"]


def test_parse_keeps_processing_when_one_attachment_is_invalid(tmp_path):
    host = Document()
    host.add_paragraph("需求名称")
    host.add_paragraph("测试需求")
    host.add_paragraph("需求内容说明")
    host.add_paragraph("功能描述")
    host.add_paragraph("主文档正文")
    host.add_paragraph("领导审核意见")

    embedded = Document()
    embedded.add_paragraph("有效附件正文")

    host_buf = BytesIO()
    embedded_buf = BytesIO()
    host.save(host_buf)
    embedded.save(embedded_buf)

    merged_bytes = _inject_docx_embedding(host_buf.getvalue(), "valid.docx", embedded_buf.getvalue())
    merged_bytes = _inject_docx_embedding(merged_bytes, "broken.docx", b"not-a-real-docx")
    file_path = tmp_path / "with-broken.docx"
    file_path.write_bytes(merged_bytes)

    result = DocxParser().parse(str(file_path))

    assert "主文档正文" in result["requirement_content"]
    assert "有效附件正文" in result["requirement_content"]


def test_parse_deduplicates_repeated_blocks_and_marks_attachment_sources(tmp_path):
    host = Document()
    host.add_paragraph("需求名称")
    host.add_paragraph("测试需求")
    host.add_paragraph("需求内容说明")
    host.add_paragraph("重复正文")
    host.add_paragraph("重复正文")
    host.add_paragraph("唯一正文")
    host.add_paragraph("领导审核意见")

    embedded = Document()
    embedded.add_paragraph("重复正文")
    embedded.add_paragraph("附件新增内容")

    host_buf = BytesIO()
    embedded_buf = BytesIO()
    host.save(host_buf)
    embedded.save(embedded_buf)

    merged_bytes = _inject_docx_embedding(host_buf.getvalue(), "attachment.docx", embedded_buf.getvalue())
    file_path = tmp_path / "dedup.docx"
    file_path.write_bytes(merged_bytes)

    result = DocxParser().parse(str(file_path))

    content = result["requirement_content"]
    assert content.count("重复正文") == 1
    assert "唯一正文" in content
    assert "附件新增内容" in content
    assert "【附件: attachment.docx】" in content
