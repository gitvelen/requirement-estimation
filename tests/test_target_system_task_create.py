import os
import sys
from datetime import datetime
from io import BytesIO

import pytest
from fastapi.testclient import TestClient
from docx import Document

ROOT_DIR = os.path.dirname(os.path.dirname(__file__))
if ROOT_DIR not in sys.path:
    sys.path.insert(0, ROOT_DIR)

from backend.api import routes as task_routes
from backend.api import system_routes
from backend.app import app
from backend.config.config import settings
from backend.service import user_service


@pytest.fixture()
def client(tmp_path, monkeypatch):
    data_dir = tmp_path / "data"
    upload_dir = tmp_path / "uploads"
    data_dir.mkdir(parents=True, exist_ok=True)
    upload_dir.mkdir(parents=True, exist_ok=True)

    monkeypatch.setattr(settings, "REPORT_DIR", str(data_dir))
    monkeypatch.setattr(settings, "UPLOAD_DIR", str(upload_dir))
    monkeypatch.setattr(settings, "DEBUG", False)

    monkeypatch.setattr(user_service, "USER_STORE_PATH", str(data_dir / "users.json"))
    monkeypatch.setattr(user_service, "USER_STORE_LOCK_PATH", str(data_dir / "users.json.lock"))

    monkeypatch.setattr(task_routes, "TASK_STORE_PATH", str(data_dir / "task_storage.json"))
    monkeypatch.setattr(task_routes, "TASK_STORE_LOCK_PATH", str(data_dir / "task_storage.json.lock"))
    monkeypatch.setattr(system_routes, "CSV_PATH", str(tmp_path / "system_list.csv"))
    monkeypatch.setattr(task_routes, "MAGIC_AVAILABLE", False)

    async def _noop_process_task(task_id: str, file_path: str):
        return None

    monkeypatch.setattr(task_routes, "process_task", _noop_process_task)
    return TestClient(app)


def _seed_user(username: str, password: str, roles):
    user = user_service.create_user_record(
        {
            "username": username,
            "display_name": username,
            "password": password,
            "roles": roles,
        }
    )
    with user_service.user_storage_context() as users:
        users.append(user)
    return user


def _login(client: TestClient, username: str, password: str) -> str:
    response = client.post("/api/v1/auth/login", json={"username": username, "password": password})
    assert response.status_code == 200
    return response.json()["data"]["token"]


def _seed_systems(systems):
    system_routes._write_systems(systems)


def _build_docx_upload():
    return (
        "requirements.docx",
        b"fake-docx-content",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _build_real_docx_upload():
    document = Document()
    document.add_paragraph("需求内容说明")
    document.add_paragraph("真实 DOCX 内容")
    buffer = BytesIO()
    document.save(buffer)
    return (
        "requirements.docx",
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _build_doc_upload():
    return (
        "requirements.doc",
        b"fake-doc-content",
        "application/msword",
    )


def test_create_task_persists_specific_target_system(client):
    manager = _seed_user("pm_target_1", "pwd123", ["manager"])
    _seed_systems(
        [
            {
                "id": "sys_pay",
                "name": "支付系统",
                "abbreviation": "PAY",
                "status": "运行中",
                "extra": {"owner_id": manager["id"], "owner_username": "pm_target_1"},
            },
            {
                "id": "sys_core",
                "name": "核心账务",
                "abbreviation": "CORE",
                "status": "运行中",
                "extra": {"owner_username": "other_pm"},
            },
        ]
    )
    token = _login(client, "pm_target_1", "pwd123")

    response = client.post(
        "/api/v1/tasks",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": _build_docx_upload()},
        data={
            "name": "目标系统任务",
            "target_system_mode": "specific",
            "target_system_name": "支付系统",
        },
    )

    assert response.status_code == 200

    with task_routes._task_storage_context() as data:
        task = next(iter(data.values()))

    assert task["name"] == "目标系统任务"
    assert task["target_system_mode"] == "specific"
    assert task["target_system_name"] == "支付系统"


def test_create_task_accepts_legacy_doc_upload(client):
    manager = _seed_user("pm_target_doc", "pwd123", ["manager"])
    _seed_systems(
        [
            {
                "id": "sys_pay",
                "name": "支付系统",
                "abbreviation": "PAY",
                "status": "运行中",
                "extra": {"owner_id": manager["id"], "owner_username": "pm_target_doc"},
            }
        ]
    )
    token = _login(client, "pm_target_doc", "pwd123")

    response = client.post(
        "/api/v1/tasks",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": _build_doc_upload()},
        data={
            "name": "老式文档任务",
            "target_system_mode": "specific",
            "target_system_name": "支付系统",
        },
    )

    assert response.status_code == 200

    with task_routes._task_storage_context() as data:
        task = next(iter(data.values()))

    assert task["filename"] == "requirements.doc"


def test_create_task_accepts_real_docx_when_magic_reports_zip(client, monkeypatch):
    monkeypatch.setattr(task_routes, "MAGIC_AVAILABLE", True)
    monkeypatch.setattr(task_routes.magic, "from_buffer", lambda *_args, **_kwargs: "application/zip")

    manager = _seed_user("pm_target_docx_zip", "pwd123", ["manager"])
    _seed_systems(
        [
            {
                "id": "sys_pay",
                "name": "支付系统",
                "abbreviation": "PAY",
                "status": "运行中",
                "extra": {"owner_id": manager["id"], "owner_username": "pm_target_docx_zip"},
            }
        ]
    )
    token = _login(client, "pm_target_docx_zip", "pwd123")

    response = client.post(
        "/api/v1/tasks",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": _build_real_docx_upload()},
        data={
            "name": "真实DOCX任务",
            "target_system_mode": "specific",
            "target_system_name": "支付系统",
        },
    )

    assert response.status_code == 200


def test_create_task_rejects_specific_system_outside_manager_scope(client):
    _seed_user("pm_target_2", "pwd123", ["manager"])
    _seed_systems(
        [
            {
                "id": "sys_core",
                "name": "核心账务",
                "abbreviation": "CORE",
                "status": "运行中",
                "extra": {"owner_username": "other_pm"},
            }
        ]
    )
    token = _login(client, "pm_target_2", "pwd123")

    response = client.post(
        "/api/v1/tasks",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": _build_docx_upload()},
        data={
            "target_system_mode": "specific",
            "target_system_name": "核心账务",
        },
    )

    assert response.status_code == 400


def test_create_task_accepts_unlimited_when_manager_has_no_owned_systems(client):
    _seed_user("pm_target_3", "pwd123", ["manager"])
    _seed_systems(
        [
            {
                "id": "sys_core",
                "name": "核心账务",
                "abbreviation": "CORE",
                "status": "运行中",
                "extra": {"owner_username": "other_pm"},
            }
        ]
    )
    token = _login(client, "pm_target_3", "pwd123")

    response = client.post(
        "/api/v1/tasks",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": _build_docx_upload()},
        data={"target_system_mode": "unlimited", "target_system_name": ""},
    )

    assert response.status_code == 200

    with task_routes._task_storage_context() as data:
        task = next(iter(data.values()))

    assert task["target_system_mode"] == "unlimited"
    assert task["target_system_name"] == ""


def test_task_detail_and_requirement_result_return_target_system_fields(client):
    manager = _seed_user("pm_target_4", "pwd123", ["manager"])
    token = _login(client, "pm_target_4", "pwd123")
    task_id = "task_target_detail"

    with task_routes._task_storage_context() as data:
        data[task_id] = {
            "task_id": task_id,
            "name": "详情任务",
            "creator_id": manager["id"],
            "creator_name": manager["display_name"],
            "status": "completed",
            "workflow_status": "draft",
            "created_at": datetime.now().isoformat(),
            "systems_data": {"支付系统": []},
            "systems": ["支付系统"],
            "modifications": [],
            "target_system_mode": "specific",
            "target_system_name": "支付系统",
        }

    detail_response = client.get(
        f"/api/v1/tasks/{task_id}",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert detail_response.status_code == 200
    detail_payload = detail_response.json()["data"]
    assert detail_payload["targetSystemMode"] == "specific"
    assert detail_payload["targetSystemName"] == "支付系统"

    result_response = client.get(
        f"/api/v1/requirement/result/{task_id}",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert result_response.status_code == 200
    result_payload = result_response.json()["data"]
    assert result_payload["target_system_mode"] == "specific"
    assert result_payload["target_system_name"] == "支付系统"
