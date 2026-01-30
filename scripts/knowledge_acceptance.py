#!/usr/bin/env python3
"""
知识库功能验收脚本

本期知识库采用本地文件向量库（无 MinIO/Milvus/Redis），本脚本测试核心能力：
1. 健康检查（后端 + 知识库）
2. 知识库导入（DOCX）
3. 知识库检索
4. 统计信息
5. 效果评估指标

使用方法：
    python scripts/knowledge_acceptance.py
"""

import os
import sys
import json
import requests
import logging
from datetime import datetime

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class KnowledgeAcceptanceTest:
    """知识库功能验收测试"""

    def __init__(self, base_url="http://127.0.0.1:443"):
        self.base_url = base_url
        self.api_base = f"{base_url}/api/v1/knowledge"
        self.admin_api_key = os.getenv("ADMIN_API_KEY", "").strip()
        self.manager_username = os.getenv("KB_USERNAME", "integration_manager").strip()
        self.manager_password = os.getenv("KB_PASSWORD", "ChangeMe123!").strip()
        self.admin_username = os.getenv("KB_ADMIN_USERNAME", "integration_admin").strip()
        self.admin_password = os.getenv("KB_ADMIN_PASSWORD", "ChangeMe123!").strip()
        self.manager_token = None
        self.admin_token = None
        self.test_results = []
        self.passed = 0
        self.failed = 0

    def log_test(self, test_name, passed, message=""):
        """记录测试结果"""
        status = "✓ PASS" if passed else "✗ FAIL"
        logger.info(f"{status} - {test_name}")
        if message:
            logger.info(f"    {message}")

        self.test_results.append({
            "test": test_name,
            "passed": passed,
            "message": message
        })

        if passed:
            self.passed += 1
        else:
            self.failed += 1

    def _request(self, method: str, path: str, *, headers=None, timeout=10, **kwargs):
        url = f"{self.base_url}{path}"
        return requests.request(method, url, headers=headers, timeout=timeout, **kwargs)

    def _login(self, username: str, password: str) -> str:
        response = self._request(
            "POST",
            "/api/v1/auth/login",
            json={"username": username, "password": password},
            timeout=10,
        )
        response.raise_for_status()
        data = response.json()
        token = (data.get("data") or {}).get("token")
        if not token:
            raise RuntimeError("登录成功但未返回token")
        return token

    def _ensure_user(self, username: str, display_name: str, roles: list, password: str) -> bool:
        headers = {}
        if self.admin_api_key:
            headers["X-API-Key"] = self.admin_api_key

        payload = {
            "username": username,
            "displayName": display_name,
            "password": password,
            "roles": roles,
            "email": f"{username}@example.com",
            "department": "knowledge_acceptance",
            "isActive": True,
        }
        response = self._request("POST", "/api/v1/users", json=payload, headers=headers, timeout=10)
        if response.status_code == 200:
            return True
        if response.status_code == 400 and "用户名已存在" in response.text:
            return True
        return False

    def _prepare_auth(self) -> bool:
        try:
            # 尽量确保测试用户存在（需要ADMIN_API_KEY或DEBUG模式放行）
            self._ensure_user(self.admin_username, "验收管理员", ["admin"], self.admin_password)
            self._ensure_user(self.manager_username, "验收项目经理", ["manager"], self.manager_password)

            self.admin_token = self._login(self.admin_username, self.admin_password)
            self.manager_token = self._login(self.manager_username, self.manager_password)

            self.log_test("准备认证", True, f"manager={self.manager_username}, admin={self.admin_username}")
            return True
        except Exception as e:
            self.log_test("准备认证", False, f"无法登录，请确认用户存在且密码正确（默认ChangeMe123!）。错误: {e}")
            return False

    def test_1_health(self):
        """测试1：健康检查"""
        logger.info("\n" + "="*60)
        logger.info("测试1：健康检查")
        logger.info("="*60)

        try:
            response = self._request("GET", "/api/v1/health", timeout=5)
            healthy = response.status_code == 200
            self.log_test("后端健康检查", healthy, f"HTTP {response.status_code}")
            return healthy

        except Exception as e:
            self.log_test("后端健康检查", False, f"测试失败: {e}")
            return False

    def test_2_knowledge_stats(self):
        """测试2：获取知识库统计信息"""
        logger.info("\n" + "="*60)
        logger.info("测试2：知识库统计信息")
        logger.info("="*60)

        try:
            response = self._request(
                "GET",
                "/api/v1/knowledge/stats",
                headers={"Authorization": f"Bearer {self.manager_token}"},
                timeout=10,
            )
            response.raise_for_status()

            data = response.json()
            stats = data.get("data", {})

            logger.info(f"知识库统计：")
            logger.info(f"  - Collection名称: {stats.get('name', 'N/A')}")
            logger.info(f"  - 知识总数: {stats.get('count', 0)} 条")
            logger.info(f"  - 索引类型: {stats.get('index', 'N/A')}")
            logger.info(f"  - 度量方式: {stats.get('metric_type', 'N/A')}")

            self.log_test(
                "获取统计信息",
                True,
                f"当前有 {stats.get('count', 0)} 条知识"
            )

            return True

        except Exception as e:
            self.log_test("获取统计信息", False, f"测试失败: {e}")
            return False

    def test_3_evaluation_metrics(self):
        """测试3：获取效果评估指标"""
        logger.info("\n" + "="*60)
        logger.info("测试3：效果评估指标")
        logger.info("="*60)

        try:
            response = self._request(
                "GET",
                "/api/v1/knowledge/evaluation-metrics",
                headers={"Authorization": f"Bearer {self.manager_token}"},
                timeout=10,
            )
            response.raise_for_status()

            data = response.json()
            metrics = data.get("data", {})

            logger.info(f"效果评估指标：")
            logger.info(f"  - 检索命中率: {metrics.get('hit_rate', 0)}%")
            logger.info(f"  - 平均相似度: {metrics.get('avg_similarity', 0)}%")
            logger.info(f"  - 案例采纳率: {metrics.get('adoption_rate', 0)}%")
            logger.info(f"  - 总检索次数: {metrics.get('total_searches', 0)}")
            logger.info(f"  - 总评估任务: {metrics.get('total_tasks', 0)}")

            quality = metrics.get('quality_comparison')
            if quality:
                logger.info(f"  - 质量对比（使用知识库）: {quality.get('with_kb', 0)}%")
                logger.info(f"  - 质量对比（未使用知识库）: {quality.get('without_kb', 0)}%")

            self.log_test(
                "获取评估指标",
                True,
                "指标获取成功"
            )

            return True

        except Exception as e:
            self.log_test("获取评估指标", False, f"测试失败: {e}")
            return False

    def test_4_import_test_data(self):
        """测试4：导入测试数据"""
        logger.info("\n" + "="*60)
        logger.info("测试4：导入测试数据")
        logger.info("="*60)

        try:
            # 生成一份可抽取的 DOCX 测试材料并上传
            from io import BytesIO
            from docx import Document

            doc = Document()
            doc.add_heading("支付中台", level=1)
            doc.add_paragraph("system_short_name：Payment")
            doc.add_paragraph("system_category：中台系统")
            doc.add_paragraph("business_goal：统一支付渠道接入，沉淀支付能力，提升接入效率与一致性。")
            doc.add_paragraph("core_functions：渠道接入、路由编排、签名验签、对账、清结算")
            doc.add_paragraph("tech_stack：Java / Spring Cloud / MySQL / Redis / Kafka")
            doc.add_paragraph("architecture：微服务 + 分布式集群；灰度发布；多活容灾")
            doc.add_paragraph("performance：TPS 10000+，关键接口RT < 200ms")
            doc.add_paragraph("main_users：各业务系统、支付运营、对账人员")
            doc.add_paragraph("notes：外部对接需符合签名规范；交易需落库并具备可追溯性。")

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)

            files = {"file": ("test_system.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            data = {
                "auto_extract": "true",
                "knowledge_type": "system_profile",
                "system_name": "支付中台",
            }

            response = self._request(
                "POST",
                "/api/v1/knowledge/import",
                headers={"Authorization": f"Bearer {self.manager_token}"},
                files=files,
                data=data,
                timeout=60,
            )

            response.raise_for_status()
            result = response.json()

            success = result.get('data', {}).get('success', 0)
            failed = result.get('data', {}).get('failed', 0)

            logger.info(f"导入结果：成功 {success} 条，失败 {failed} 条")

            if success > 0:
                self.log_test(
                    "导入测试数据",
                    True,
                    f"成功导入 {success} 条系统知识"
                )
                return True

            self.log_test(
                "导入测试数据",
                False,
                "没有数据导入成功"
            )
            return False

        except Exception as e:
            self.log_test("导入测试数据", False, f"测试失败: {e}")
            return False

    def test_5_knowledge_search(self):
        """测试5：知识检索测试"""
        logger.info("\n" + "="*60)
        logger.info("测试5：知识检索测试")
        logger.info("="*60)

        try:
            # 测试检索
            search_data = {
                "query": "支付中台 微信支付",
                "system_name": "支付中台",
                "top_k": 5,
                # 由于embedding/阈值会受模型与语料影响，验收时用较低阈值保证可观测性
                "similarity_threshold": 0.2
            }

            response = self._request(
                "POST",
                "/api/v1/knowledge/search",
                headers={"Authorization": f"Bearer {self.manager_token}"},
                json=search_data,
                timeout=20,
            )

            response.raise_for_status()
            result = response.json()

            total = result.get('data', {}).get('total', 0)
            results = result.get('data', {}).get('results', [])

            logger.info(f"检索结果：找到 {total} 条相关知识")

            if total > 0:
                logger.info(f"\n最相关的知识：")
                for i, item in enumerate(results[:3], 1):
                    logger.info(f"  {i}. {item.get('system_name', 'N/A')} - 相似度: {item.get('similarity', 0):.2f}")
                    logger.info(f"     {item.get('content', '')[:100]}...")

                self.log_test(
                    "知识检索",
                    True,
                    f"成功检索到 {total} 条相关知识"
                )
                return True
            else:
                self.log_test(
                    "知识检索",
                    False,
                    "未检索到相关知识（可能是知识库为空）"
                )
                return False

        except Exception as e:
            self.log_test("知识检索", False, f"测试失败: {e}")
            return False

    def test_6_rebuild_index(self):
        """测试6：重建索引（本地向量库）"""
        logger.info("\n" + "="*60)
        logger.info("测试6：重建索引")
        logger.info("="*60)

        try:
            response = self._request(
                "POST",
                "/api/v1/knowledge/rebuild-index",
                headers={"Authorization": f"Bearer {self.admin_token}"},
                timeout=30,
            )
            response.raise_for_status()
            result = response.json()
            data = result.get("data") or {}
            ok = result.get("code") == 200 and (data.get("status") in ("success", None) or data.get("message"))
            self.log_test("重建索引", ok, data.get("message", ""))
            return ok
        except Exception as e:
            self.log_test("重建索引", False, f"测试失败: {e}")
            return False

    def run_all_tests(self):
        """运行所有测试"""
        logger.info("\n" + "="*60)
        logger.info("知识库功能验收测试")
        logger.info("="*60)
        logger.info(f"测试时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        logger.info(f"API地址: {self.api_base}")
        logger.info("="*60)

        # 认证准备（后续接口需要登录）
        if not self._prepare_auth():
            self.print_summary()
            self.generate_report()
            return

        # 运行测试
        tests = [
            self.test_1_health,
            self.test_2_knowledge_stats,
            self.test_3_evaluation_metrics,
            self.test_4_import_test_data,
            self.test_5_knowledge_search,
            self.test_6_rebuild_index,
        ]

        for test in tests:
            try:
                test()
            except Exception as e:
                logger.error(f"测试执行异常: {e}")

        # 输出测试总结
        self.print_summary()

        # 生成验收报告
        self.generate_report()

    def print_summary(self):
        """打印测试总结"""
        logger.info("\n" + "="*60)
        logger.info("测试总结")
        logger.info("="*60)
        logger.info(f"总测试数: {len(self.test_results)}")
        logger.info(f"通过: {self.passed}")
        logger.info(f"失败: {self.failed}")
        logger.info(f"通过率: {self.passed / len(self.test_results) * 100:.1f}%")

        if self.failed == 0:
            logger.info("\n✓ 所有测试通过！知识库功能验收合格。")
        else:
            logger.info(f"\n✗ 有 {self.failed} 个测试失败，请检查日志。")

    def generate_report(self):
        """生成验收报告"""
        report = {
            "test_time": datetime.now().isoformat(),
            "total_tests": len(self.test_results),
            "passed": self.passed,
            "failed": self.failed,
            "pass_rate": f"{self.passed / len(self.test_results) * 100:.1f}%",
            "results": self.test_results
        }

        # 保存到文件
        report_file = "data/knowledge_acceptance_report.json"
        os.makedirs("data", exist_ok=True)

        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

        logger.info(f"\n验收报告已保存到: {report_file}")


def main():
    """主函数"""
    import argparse

    parser = argparse.ArgumentParser(description="知识库功能验收测试")
    parser.add_argument(
        "--base-url",
        default="http://127.0.0.1:443",
        help="API基础地址（默认: http://127.0.0.1:443）"
    )

    args = parser.parse_args()

    # 创建测试实例
    tester = KnowledgeAcceptanceTest(base_url=args.base_url)

    # 运行所有测试
    tester.run_all_tests()

    # 根据测试结果设置退出码
    sys.exit(0 if tester.failed == 0 else 1)


if __name__ == "__main__":
    main()
