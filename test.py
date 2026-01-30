"""
系统测试脚本
测试知识库导入和AI评估功能
"""
import requests
import time
import json

BASE_URL = "http://localhost:443/api/v1"

def login(username: str, password: str) -> str:
    """登录并返回JWT Token"""
    resp = requests.post(
        f"{BASE_URL}/auth/login",
        json={"username": username, "password": password},
        timeout=20
    )
    resp.raise_for_status()
    data = resp.json().get("data") or {}
    token = data.get("token") or ""
    if not token:
        raise RuntimeError("登录成功但未返回token")
    return token

def test_health():
    """测试健康检查"""
    print("\n=== 测试1: 健康检查 ===")
    try:
        response = requests.get(f"{BASE_URL}/health")
        print(f"状态码: {response.status_code}")
        print(f"响应: {response.json()}")
        return response.status_code == 200
    except Exception as e:
        print(f"❌ 失败: {e}")
        return False

def test_knowledge_health(token: str):
    """测试知识库健康检查"""
    print("\n=== 测试2: 知识库健康检查 ===")
    try:
        response = requests.get(
            f"{BASE_URL}/knowledge/health",
            headers={"Authorization": f"Bearer {token}"},
            timeout=20
        )
        print(f"状态码: {response.status_code}")
        print(f"响应: {response.json()}")
        return response.status_code == 200
    except Exception as e:
        print(f"❌ 失败: {e}")
        return False

def test_import_knowledge(token: str):
    """测试导入知识库"""
    print("\n=== 测试3: 导入系统知识 ===")
    try:
        from io import BytesIO
        from docx import Document

        doc = Document()
        doc.add_heading("支付中台", level=1)
        doc.add_paragraph("system_short_name：Payment")
        doc.add_paragraph("business_goal：统一支付渠道接入，沉淀支付能力。")
        doc.add_paragraph("core_functions：渠道接入、路由编排、签名验签、对账、清结算")
        doc.add_paragraph("tech_stack：Java / Spring Cloud / MySQL / Redis / Kafka")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        files = {"file": ("test_system.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"auto_extract": "true", "knowledge_type": "system_profile", "system_name": "支付中台"}
        response = requests.post(
            f"{BASE_URL}/knowledge/import",
            headers={"Authorization": f"Bearer {token}"},
            files=files,
            data=data,
            timeout=60
        )
        
        print(f"状态码: {response.status_code}")
        result = response.json()
        print(f"响应: {json.dumps(result, indent=2, ensure_ascii=False)}")
        return response.status_code == 200
    except Exception as e:
        print(f"❌ 失败: {e}")
        return False

def test_knowledge_stats(token: str):
    """测试知识库统计"""
    print("\n=== 测试4: 知识库统计 ===")
    try:
        response = requests.get(
            f"{BASE_URL}/knowledge/stats",
            headers={"Authorization": f"Bearer {token}"},
            params={"system_name": "支付中台"},
            timeout=20
        )
        print(f"状态码: {response.status_code}")
        result = response.json()
        print(f"响应: {json.dumps(result, indent=2, ensure_ascii=False)}")
        return response.status_code == 200
    except Exception as e:
        print(f"❌ 失败: {e}")
        return False

def test_search_knowledge(token: str):
    """测试知识检索"""
    print("\n=== 测试5: 知识检索 ===")
    try:
        payload = {
            "query": "账户管理功能",
            "system_name": "支付中台",
            "top_k": 3,
            "similarity_threshold": 0.6
        }
        response = requests.post(
            f"{BASE_URL}/knowledge/search",
            headers={"Authorization": f"Bearer {token}"},
            json=payload,
            timeout=30
        )
        print(f"状态码: {response.status_code}")
        result = response.json()
        print(f"响应: {json.dumps(result, indent=2, ensure_ascii=False)}")
        return response.status_code == 200
    except Exception as e:
        print(f"❌ 失败: {e}")
        return False

def main():
    """运行所有测试"""
    print("╔══════════════════════════════════════════╗")
    print("║   业务需求工作量评估系统 - 自动化测试    ║")
    print("╚══════════════════════════════════════════╝")
    
    username = "integration_manager"
    password = "ChangeMe123!"
    try:
        token = login(username, password)
    except Exception as e:
        print(f"❌ 登录失败: {e}")
        return

    tests = [
        ("健康检查", test_health),
        ("知识库健康检查", lambda: test_knowledge_health(token)),
        ("导入系统知识", lambda: test_import_knowledge(token)),
        ("知识库统计", lambda: test_knowledge_stats(token)),
        ("知识检索", lambda: test_search_knowledge(token)),
    ]
    
    results = []
    for name, test_func in tests:
        try:
            success = test_func()
            results.append((name, success))
            time.sleep(1)  # 避免请求过快
        except Exception as e:
            print(f"❌ {name} 测试异常: {e}")
            results.append((name, False))
    
    # 打印测试结果
    print("\n" + "="*50)
    print("测试结果汇总:")
    print("="*50)
    
    passed = 0
    for name, success in results:
        status = "✅ 通过" if success else "❌ 失败"
        print(f"{status} - {name}")
        if success:
            passed += 1
    
    print("="*50)
    print(f"总计: {passed}/{len(results)} 通过")
    print("="*50)

if __name__ == "__main__":
    main()
