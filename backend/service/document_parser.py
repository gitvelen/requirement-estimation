"""
多格式文档解析服务
支持CSV、DOCX、XLSX、PDF等多种格式的文档解析
"""
import logging
import csv
import json
from io import StringIO, BytesIO
from typing import List, Dict, Any, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentParser:
    """多格式文档解析器"""

    # 支持的文件类型
    SUPPORTED_TYPES = {
        "csv": "text/csv",
        "doc": "application/msword",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xls": "application/vnd.ms-excel",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "pdf": "application/pdf",
        "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    }

    def __init__(self):
        """初始化解析器"""
        logger.info("文档解析服务初始化完成")

    def parse(
        self,
        file_content: bytes,
        filename: str,
        file_type: str = None
    ) -> List[Dict[str, Any]]:
        """
        解析文档

        Args:
            file_content: 文件内容（字节）
            filename: 文件名
            file_type: 文件类型（如果为None则从filename推断）

        Returns:
            list: 解析后的数据列表

        Raises:
            ValueError: 不支持的文件类型
            Exception: 解析失败
        """
        # 推断文件类型
        if file_type is None:
            file_ext = Path(filename).suffix.lower().lstrip('.')
            file_type = file_ext

        # 验证文件类型
        if file_type not in self.SUPPORTED_TYPES:
            raise ValueError(
                f"不支持的文件类型: {file_type}，"
                f"支持的类型: {', '.join(self.SUPPORTED_TYPES.keys())}"
            )

        logger.info(f"开始解析文档: {filename} (类型: {file_type})")

        # 根据类型解析
        if file_type == "csv":
            return self._parse_csv(file_content)
        elif file_type == "doc":
            return self._parse_doc(file_content, filename)
        elif file_type == "docx":
            return self._parse_docx(file_content)
        elif file_type == "xls":
            return self._parse_xls(file_content, filename)
        elif file_type == "xlsx":
            return self._parse_xlsx(file_content)
        elif file_type == "pdf":
            return self._parse_pdf(file_content)
        elif file_type == "pptx":
            return self._parse_pptx(file_content)
        else:
            raise ValueError(f"未实现的文件类型: {file_type}")

    def _parse_csv(self, file_content: bytes) -> List[Dict[str, Any]]:
        """
        解析CSV文件

        支持的CSV格式：
        - 系统知识库（system_profile）
        """
        try:
            # 解码为文本
            text = file_content.decode('utf-8-sig')  # 处理BOM

            # 使用CSV DictReader解析
            reader = csv.DictReader(StringIO(text))

            # 转换为字典列表
            data = []
            for row in reader:
                # 过滤空行
                if any(row.values()):
                    data.append(dict(row))

            logger.info(f"CSV解析成功，共 {len(data)} 行")

            return data

        except UnicodeDecodeError:
            # 尝试GBK编码
            try:
                text = file_content.decode('gbk')
                reader = csv.DictReader(StringIO(text))
                data = [dict(row) for row in reader if any(row.values())]
                logger.info(f"CSV解析成功（GBK编码），共 {len(data)} 行")
                return data
            except Exception as e:
                logger.error(f"CSV解析失败（编码错误）: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"CSV解析失败: {str(e)}")
            raise

    def _parse_docx(self, file_content: bytes) -> List[Dict[str, Any]]:
        """
        解析DOCX文件

        提取内容：
        - 段落文本
        - 表格数据
        - 标题层级
        """
        try:
            from docx import Document
            from docx.table import Table

            # 从字节加载文档
            doc = Document(BytesIO(file_content))

            data = []

            # 提取段落
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "type": "paragraph",
                        "text": para.text.strip(),
                        "style": para.style.name if para.style else "Normal"
                    })

            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "type": "table",
                    "data": table_data
                })

            logger.info(f"DOCX解析成功: {len(paragraphs)}个段落, {len(tables)}个表格")

            # 返回结构化数据
            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": {
                    "total_paragraphs": len(paragraphs),
                    "total_tables": len(tables)
                }
            }

        except ImportError:
            logger.error("未安装python-docx库，请运行: pip install python-docx")
            raise Exception("缺少DOCX解析库")
        except Exception as e:
            logger.error(f"DOCX解析失败: {str(e)}")
            raise

    def _parse_xlsx(self, file_content: bytes) -> List[Dict[str, Any]]:
        """
        解析XLSX文件

        支持多个Sheet（解析为表格数据结构，结构化抽取以 system_profile 为主）
        """
        try:
            from openpyxl import load_workbook

            # 从字节加载工作簿，data_only=True 读取公式计算后的值
            wb = load_workbook(BytesIO(file_content), data_only=True)

            data = {}

            # ESB服务治理文档需要忽略的sheet
            ignored_sheets = {"系统清单", "字典", "新服务治理平台服务视图"}

            # 遍历所有Sheet
            logger.info(f"文件包含的所有sheets: {wb.sheetnames}")
            for sheet_name in wb.sheetnames:
                # 跳过需要忽略的sheet
                if sheet_name in ignored_sheets:
                    logger.info(f"✓ 跳过Sheet '{sheet_name}'（已配置忽略）")
                    continue
                logger.info(f"✗ 处理Sheet '{sheet_name}'（未在忽略列表中）")

                ws = wb[sheet_name]

                # 保留原始行结构，具体表头识别交给调用方按业务场景处理
                sheet_data = []
                for row in ws.iter_rows(values_only=True):
                    if any(cell is not None and str(cell).strip() != "" for cell in row):
                        sheet_data.append(list(row))

                data[sheet_name] = sheet_data
                logger.info(f"Sheet '{sheet_name}': {len(sheet_data)} 行数据")

            logger.info(f"XLSX解析成功，共 {len(data)} 个Sheet")

            return data

        except ImportError:
            logger.error("未安装openpyxl库，请运行: pip install openpyxl")
            raise Exception("缺少XLSX解析库")
        except Exception as e:
            logger.error(f"XLSX解析失败: {str(e)}")
            raise

    def _parse_pdf(self, file_content: bytes) -> List[Dict[str, Any]]:
        """
        解析PDF文件（基础功能）

        提取：
        - 文本内容
        - 元数据（作者、创建时间等）

        Note: PDF解析较为复杂，此为基础实现
        """
        try:
            import PyPDF2

            # 从字节加载PDF
            pdf_file = BytesIO(file_content)
            reader = PyPDF2.PdfReader(pdf_file)

            # 提取文本
            text_content = []
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append({
                            "page": page_num + 1,
                            "text": text.strip()
                        })
                except Exception as e:
                    logger.warning(f"第{page_num + 1}页解析失败: {e}")

            # 提取元数据
            metadata = reader.metadata

            logger.info(f"PDF解析成功: {len(text_content)}页")

            return {
                "pages": text_content,
                "metadata": {
                    "author": metadata.get("/Author", ""),
                    "creator": metadata.get("/Creator", ""),
                    "title": metadata.get("/Title", ""),
                    "page_count": len(reader.pages)
                }
            }

        except ImportError:
            logger.warning("未安装PyPDF2库，PDF解析功能不可用")
            logger.warning("请运行: pip install PyPDF2")
            return {"pages": [], "metadata": {}, "error": "缺少PDF解析库"}
        except Exception as e:
            logger.error(f"PDF解析失败: {str(e)}")
            raise

    def _parse_pptx(self, file_content: bytes) -> List[Dict[str, Any]]:
        """
        解析PPTX文件

        提取：
        - 幻灯片文本
        - 标题层级
        """
        try:
            import re
            import zipfile
            from xml.etree import ElementTree as ET

            # PPTX 本质是 zip + XML，直接解析 slide XML 中的 <a:t> 文本即可。
            slides_data = []
            with zipfile.ZipFile(BytesIO(file_content)) as zf:
                slide_files = [
                    name
                    for name in zf.namelist()
                    if name.startswith("ppt/slides/slide") and name.endswith(".xml")
                ]

                def slide_key(name: str) -> int:
                    match = re.search(r"slide(\d+)\.xml$", name)
                    return int(match.group(1)) if match else 0

                slide_files.sort(key=slide_key)

                for slide_path in slide_files:
                    try:
                        xml_bytes = zf.read(slide_path)
                        root = ET.fromstring(xml_bytes)
                        texts = []
                        for node in root.iter():
                            tag = node.tag or ""
                            if tag.endswith("}t") or tag == "a:t":
                                if node.text and node.text.strip():
                                    texts.append(node.text.strip())
                        slide_text = "\n".join(texts).strip()
                        if slide_text:
                            slides_data.append(
                                {
                                    "slide": slide_key(slide_path) or (len(slides_data) + 1),
                                    "text": slide_text,
                                }
                            )
                    except Exception as e:
                        logger.warning(f"PPTX幻灯片解析失败: {slide_path}: {e}")
                        continue

                total_slides = len(slide_files)

            logger.info(f"PPTX解析成功: {len(slides_data)}张幻灯片")

            return {
                "slides": slides_data,
                "metadata": {
                    "total_slides": total_slides
                }
            }
        except Exception as e:
            logger.error(f"PPTX解析失败: {str(e)}")
            raise

    def _parse_doc(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        解析 DOC（Word 97-2003）文件

        通过 headless libreoffice 转换为纯文本（满足 REQ-NF-007：隔离目录/超时/清理）。
        """
        from backend.utils.old_format_parser import doc_bytes_to_text

        text = doc_bytes_to_text(file_content, filename)
        return {"text": text}

    def _parse_xls(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        解析 XLS（Excel 97-2003）文件

        通过 headless libreoffice 转换为 xlsx 后复用 openpyxl 解析（满足 REQ-NF-007：隔离目录/超时/清理）。
        """
        from backend.utils.old_format_parser import xls_bytes_to_sheet_rows

        return xls_bytes_to_sheet_rows(file_content, filename)

    def extract_system_knowledge(self, parsed_data: Any, expected_type: Optional[str] = None) -> Dict[str, Any]:
        """
        从解析后的数据中提取系统知识

        Args:
            parsed_data: 解析后的数据（来自parse方法）

        Returns:
            dict: 系统知识结构化数据
        """
        try:
            normalized_expected = expected_type if expected_type in ("system_profile",) else None

            # 如果是CSV数据（系统知识库格式）
            if isinstance(parsed_data, list) and len(parsed_data) > 0:
                first_row = parsed_data[0]

                # 如果指定了期望类型，优先按期望类型解析
                if normalized_expected == "system_profile":
                    extracted = self._extract_system_profile(parsed_data)
                    if extracted.get("count", 0) > 0:
                        return extracted

                # 识别字段
                if "系统名称" in first_row or "系统简称" in first_row:
                    # 系统知识库格式
                    return self._extract_system_profile(parsed_data)

            # 如果是DOCX数据
            if isinstance(parsed_data, dict) and "paragraphs" in parsed_data:
                return self._extract_from_docx(parsed_data, normalized_expected)

            # 如果是PPTX数据
            if isinstance(parsed_data, dict) and "slides" in parsed_data:
                return self._extract_from_pptx(parsed_data)

            # 如果是XLSX数据
            if isinstance(parsed_data, dict) and "metadata" not in parsed_data:
                return self._extract_from_xlsx(parsed_data, normalized_expected)

            logger.warning("无法识别的知识库格式")
            return {}

        except Exception as e:
            logger.error(f"提取系统知识失败: {str(e)}")
            return {}

    def _extract_system_profile(self, data: List[Dict]) -> Dict[str, Any]:
        """从CSV数据提取系统知识"""
        systems = []

        for row in data:
            system = {
                "system_name": row.get("系统名称", ""),
                "system_short_name": row.get("系统简称", ""),
                "system_category": row.get("系统分类", ""),
                "business_goal": row.get("业务目标", ""),
                "core_functions": row.get("核心功能", ""),
                "tech_stack": row.get("技术栈", ""),
                "architecture": row.get("架构特点", ""),
                "performance": row.get("性能指标", ""),
                "main_users": row.get("主要用户", ""),
                "notes": row.get("备注", ""),
                # 【新增】校准字段（可选）
                "in_scope": row.get("系统职责/边界", "") or row.get("系统边界", "") or row.get("in_scope", ""),
                "out_of_scope": row.get("系统不做什么", "") or row.get("out_of_scope", ""),
                "integration_points": row.get("主要集成点/上下游", "") or row.get("集成点/上下游", "") or row.get("integration_points", ""),
                "key_constraints": row.get("关键约束", "") or row.get("key_constraints", ""),
            }

            # 过滤空记录
            if system["system_name"]:
                systems.append(system)

        return {
            "type": "system_profile",
            "count": len(systems),
            "systems": systems
        }

    def _extract_from_docx(self, data: Dict, expected_type: Optional[str] = None) -> Dict[str, Any]:
        """从DOCX数据提取知识"""
        try:
            # 尝试智能提取
            extracted = self._intelligent_extract_docx(data)

            if extracted and extracted.get("systems"):
                return {
                    "type": "system_profile",
                    "count": len(extracted["systems"]),
                    "systems": extracted["systems"]
                }

            return {}

        except Exception as e:
            logger.warning(f"智能提取失败: {e}")
            return {}

    def _intelligent_extract_system_profile_from_text(self, full_text: str) -> Dict[str, Any]:
        """使用LLM从纯文本中提取系统知识(system_profile)。"""
        try:
            from backend.utils.llm_client import llm_client

            prompt = f"""请从以下系统介绍/系统架构材料中，提取“系统知识(system_profile)”的结构化信息，并按JSON返回。

材料内容：
{(full_text or '')[:3500]}

请以JSON格式返回（只返回JSON，不要任何解释）：
{{
  "systems": [
    {{
      "system_name": "系统名称",
      "system_short_name": "系统简称",
      "system_category": "系统分类",
      "business_goal": "业务目标",
      "core_functions": "核心功能（用顿号分隔）",
      "tech_stack": "技术栈",
      "architecture": "架构特点",
      "performance": "性能指标",
      "main_users": "主要用户",
      "notes": "备注",
      "in_scope": "系统职责/边界（做什么）",
      "out_of_scope": "系统明确不做什么（避免误拆）",
      "integration_points": "主要集成点/上下游（用顿号分隔）",
      "key_constraints": "关键约束（合规/性能/数据敏感/发布窗口等）"
    }}
  ]
}}
"""

            response = llm_client.chat_with_system_prompt(
                system_prompt="你是一个专业的技术文档分析助手，擅长从系统介绍/架构材料中提取结构化信息。",
                user_prompt=prompt,
                temperature=0.2,
                max_tokens=2000,
            )
            result = llm_client.extract_json(response)
            systems = result.get("systems") if isinstance(result, dict) else None
            if isinstance(systems, list) and systems:
                logger.info(f"LLM智能提取完成: systems={len(systems)}")
                return {"systems": systems}
            return {}
        except Exception as e:
            logger.warning(f"LLM智能提取失败: {e}")
            return {}

    def _intelligent_extract_docx(self, data: Dict) -> Dict[str, Any]:
        """
        使用LLM智能提取DOCX中的结构化知识

        Args:
            data: DOCX解析后的数据

        Returns:
            dict: 提取的结构化数据
        """
        try:
            # 提取所有文本内容
            all_text = []
            if "paragraphs" in data:
                for para in data["paragraphs"]:
                    all_text.append(para["text"])

            if "tables" in data:
                for table in data["tables"]:
                    for row in table["data"]:
                        all_text.append(" | ".join(row))

            full_text = "\n".join(all_text)
            return self._intelligent_extract_system_profile_from_text(full_text)
        except Exception as e:
            logger.warning(f"DOCX智能提取失败: {e}")
            return {}

    def _extract_from_pptx(self, data: Dict) -> Dict[str, Any]:
        """从PPTX数据提取系统知识"""
        try:
            slide_lines = []
            for slide in data.get("slides") or []:
                if not isinstance(slide, dict):
                    continue
                idx = slide.get("slide")
                text = str(slide.get("text") or "").strip()
                if not text:
                    continue
                slide_lines.append(f"[Slide {idx}] {text}")

            full_text = "\n\n".join(slide_lines)
            extracted = self._intelligent_extract_system_profile_from_text(full_text)
            if extracted and extracted.get("systems"):
                return {
                    "type": "system_profile",
                    "count": len(extracted["systems"]),
                    "systems": extracted["systems"],
                }
            return {}
        except Exception as e:
            logger.warning(f"PPTX智能提取失败: {e}")
            return {}

    def _extract_from_xlsx(self, data: Dict, expected_type: Optional[str] = None) -> Dict[str, Any]:
        """从XLSX数据提取知识"""
        # ESB服务治理文档需要忽略的sheet（与 _parse_xlsx 保持一致）
        ignored_sheets = {"系统清单", "字典", "新服务治理平台服务视图"}

        # 查找系统知识相关的Sheet
        for sheet_name, sheet_data in data.items():
            # 跳过需要忽略的sheet
            if sheet_name in ignored_sheets:
                logger.info(f"提取阶段跳过Sheet '{sheet_name}'（已配置忽略）")
                continue

            if len(sheet_data) > 0:
                first_row = sheet_data[0]

                normalized_expected = expected_type if expected_type in ("system_profile",) else None
                if normalized_expected:
                    headers = [str(cell) for cell in first_row]
                    rows = []
                    for row in sheet_data[1:]:
                        row_dict = {}
                        for i, cell in enumerate(row):
                            if i < len(headers):
                                row_dict[headers[i]] = str(cell) if cell is not None else ""
                        rows.append(row_dict)

                    extracted = self._extract_system_profile(rows)
                    if extracted.get("count", 0) > 0:
                        return extracted

                # 判断类型
                if "系统名称" in first_row or "系统简称" in first_row:
                    # 转换为字典列表
                    headers = [str(cell) for cell in first_row]
                    rows = []
                    for row in sheet_data[1:]:
                        row_dict = {}
                        for i, cell in enumerate(row):
                            if i < len(headers):
                                row_dict[headers[i]] = str(cell) if cell is not None else ""
                        rows.append(row_dict)

                    return self._extract_system_profile(rows)

        return {}


# 全局解析器实例
_document_parser = None


def get_document_parser() -> DocumentParser:
    """获取文档解析器单例"""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
    return _document_parser
