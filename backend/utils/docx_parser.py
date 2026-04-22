"""
文档解析模块
负责解析.docx文件，提取需求内容说明
"""
import logging
from typing import Optional, Dict, List
from docx import Document
from backend.config.config import settings
from backend.service.document_parser import DocumentParser

logger = logging.getLogger(__name__)


class DocxParser:
    """Word文档解析器"""

    def __init__(self):
        """初始化解析器"""
        self.target_section = "需求内容说明"
        self.document_parser = DocumentParser()
        logger.info("DocxParser初始化完成")

    def parse(self, file_path: str) -> Dict[str, str]:
        """
        解析Word文档，提取关键信息

        Args:
            file_path: 文档文件路径

        Returns:
            Dict: 包含需求名称、需求简述、需求内容说明等信息的字典

        Raises:
            FileNotFoundError: 文件不存在
            Exception: 解析失败
        """
        try:
            doc = Document(file_path)

            # 提取所有段落文本
            paragraphs = [para.text.strip() for para in doc.paragraphs]

            # 提取表格数据
            tables_data = self._extract_tables(doc)

            # 提取需求基本信息
            requirement_info = self._extract_basic_info(paragraphs, tables_data)

            # 提取需求内容说明
            content_section = self._extract_content_section(paragraphs)

            # 如果没有找到需求内容说明，尝试从表格中提取
            if not content_section or len(content_section) < 10:
                content_section = self._extract_content_from_tables(tables_data)

            merged_content = self._merge_attachment_content(file_path, content_section)

            result = {
                "requirement_name": requirement_info.get("requirement_name", ""),
                "requirement_summary": requirement_info.get("requirement_summary", ""),
                "requirement_content": merged_content,
                "basic_info": requirement_info,
                "all_paragraphs": paragraphs
            }

            logger.info(f"文档解析完成: {file_path}")
            logger.info(f"提取段落数: {len(paragraphs)}")
            logger.info(f"需求内容说明长度: {len(merged_content)}")

            return result

        except FileNotFoundError:
            logger.error(f"文件不存在: {file_path}")
            raise
        except Exception as e:
            logger.error(f"文档解析失败: {str(e)}")
            raise

    def _extract_tables(self, doc: Document) -> List[List[List[str]]]:
        """
        提取文档中的所有表格

        Args:
            doc: Document对象

        Returns:
            List: 三维列表，表示所有表格的行列数据
        """
        tables = []

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        logger.debug(f"提取表格数量: {len(tables)}")
        return tables

    def _extract_basic_info(
        self,
        paragraphs: List[str],
        tables: List[List[List[str]]]
    ) -> Dict[str, str]:
        """
        从段落和表格中提取需求基本信息

        Args:
            paragraphs: 段落列表
            tables: 表格列表

        Returns:
            Dict: 包含需求名称、简述等信息的字典
        """
        info = {}

        # 从段落中查找关键信息
        for i, para in enumerate(paragraphs):
            if "需求名称" in para and i + 1 < len(paragraphs):
                info["requirement_name"] = paragraphs[i + 1]
            elif "需求简述" in para or "需求概述" in para:
                # 提取后续几段作为简述
                summary_parts = []
                for j in range(i + 1, min(i + 5, len(paragraphs))):
                    if paragraphs[j] and not paragraphs[j].startswith("一、") and not paragraphs[j].startswith("1."):
                        summary_parts.append(paragraphs[j])
                    else:
                        break
                info["requirement_summary"] = "\n".join(summary_parts)

        # 从表格中提取基本信息（如果有表格）
        if tables:
            # 假设第一个表格是基本信息表格
            first_table = tables[0]
            for row in first_table:
                if len(row) >= 2:
                    key = row[0]
                    value = self._pick_basic_info_value(key, row)
                    if "需求名称" in key:
                        info["requirement_name"] = value
                    elif "需求简述" in key or "需求概述" in key:
                        info["requirement_summary"] = value

        return info

    def _pick_basic_info_value(self, key: str, row: List[str]) -> str:
        """
        从基本信息表格行中选择真实值。

        部分模板在合并单元格场景下会把表头文本重复写进多个单元格，
        这里优先取第一个“非空且不等于表头”的单元格值。
        """
        normalized_key = str(key or "").strip()
        if len(row) <= 1:
            return ""

        for cell in row[1:]:
            value = str(cell or "").strip()
            if value and value != normalized_key:
                return value

        return str(row[1] or "").strip()

    def _extract_content_section(self, paragraphs: List[str]) -> str:
        """
        提取需求内容说明章节

        Args:
            paragraphs: 段落列表

        Returns:
            str: 需求内容说明文本
        """
        content_parts = []
        found_section = False
        section_end_markers = ["领导审核意见", "相关业务材料", "附件", "审批流程"]

        for para in paragraphs:
            # 查找需求内容说明章节
            if self.target_section in para or "需求功能要点描述" in para:
                found_section = True
                continue

            # 如果找到了章节，开始收集内容
            if found_section:
                # 检查是否到达章节结尾
                if any(marker in para for marker in section_end_markers):
                    break

                # 过滤空行和标题行
                if para and not para.startswith("附件") and not para.startswith("审批"):
                    content_parts.append(para)

        content = "\n".join(content_parts)

        # 如果没有找到需求内容说明章节，尝试从表格中提取
        if not content or len(content) < 10:
            logger.info("未找到需求内容说明章节，尝试从表格提取")

        logger.debug(f"需求内容说明提取完成，长度: {len(content)}")
        return content

    def _extract_content_from_tables(self, tables: List[List[List[str]]]) -> str:
        """
        从表格中提取需求内容

        Args:
            tables: 表格列表

        Returns:
            str: 需求内容文本
        """
        content_parts = []

        for table in tables:
            for row in table:
                if len(row) >= 2:
                    key = row[0] if row[0] else ""
                    value = row[2] if len(row) > 2 else (row[1] if len(row) > 1 else "")

                    # 提取需求起因、目的等关键信息
                    if any(keyword in key for keyword in ["需求起因", "需求目的", "需求简述", "需求概述", "功能描述", "业务背景"]):
                        if value and value.strip():
                            content_parts.append(f"{key}: {value}")

        content = "\n".join(content_parts)
        logger.info(f"从表格提取需求内容，长度: {len(content)}")
        return content

    def _merge_attachment_content(self, file_path: str, base_content: str) -> str:
        try:
            with open(file_path, "rb") as fh:
                parsed = self.document_parser.parse(fh.read(), filename=file_path, file_type="docx")
        except Exception as exc:
            logger.warning("读取嵌入附件失败，回退主文档正文: %s", exc)
            return base_content

        parts: List[str] = []
        seen = set()

        def append_block(text: str) -> None:
            normalized = str(text or "").strip()
            if not normalized or normalized in seen:
                return
            seen.add(normalized)
            parts.append(normalized)

        for line in str(base_content or "").splitlines():
            append_block(line)

        for item in parsed.get("attachments") or []:
            if not isinstance(item, dict):
                continue
            attachment_name = str(item.get("name") or "").strip()
            attachment_text = str(item.get("text") or "").strip()
            if not attachment_text:
                continue
            if attachment_name:
                append_block(f"【附件: {attachment_name}】")
            for line in attachment_text.splitlines():
                append_block(line)

        return "\n".join(parts)

    def validate_file(self, file_path: str) -> bool:
        """
        验证文件是否有效

        Args:
            file_path: 文件路径

        Returns:
            bool: 文件是否有效
        """
        import os

        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return False

        # 检查文件扩展名
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in settings.ALLOWED_EXTENSIONS:
            logger.error(f"不支持的文件类型: {file_ext}")
            return False

        # 检查文件大小
        file_size = os.path.getsize(file_path)
        if file_size > settings.MAX_FILE_SIZE:
            logger.error(f"文件过大: {file_size} > {settings.MAX_FILE_SIZE}")
            return False

        logger.info(f"文件验证通过: {file_path}")
        return True


# 全局解析器实例
docx_parser = DocxParser()
